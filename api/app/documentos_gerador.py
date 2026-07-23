"""Geração dos documentos de um projeto (Memorial, Contrato, Procuração, RT).

Porta pro Python o que as macros do `CONTROLE PROJETOS ENGENHARIA.xlsm` faziam
no Word: abrir o modelo .docx e trocar as variáveis `#NOME PF`, `#CPF`, ... pelo
valor do projeto. Os modelos ficam em `api/templates_documentos/` (um .docx por
documento) — são os mesmos que a engenharia já usa, só precisam ser copiados
pra lá com as variáveis no corpo.

O DWG NÃO é gerado aqui: as macros usam automação COM do AutoCAD
(`AcadDoc`/`AlterarTextos`), que exige uma máquina Windows com AutoCAD
instalado. Ver `dwg.py` (pendente) e o contrato 11_geracao_documentos.md.
"""

import io
import logging
import re
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

DOCUMENTOS = ("memorial", "contrato", "procuracao", "rt")

MODELOS = {
    "memorial": "memorial.docx",
    "contrato": "contrato.docx",
    "procuracao": "procuracao.docx",
    "rt": "rt.docx",
}

# Prefixo do arquivo final, no padrão que a engenharia já usa nas pastas
# (ex: MD-2026-1035-750-001-A - Memorial Descritivo ...).
PREFIXOS = {"memorial": "MD", "contrato": "CT", "procuracao": "DC", "rt": "RT"}


class GeracaoError(Exception):
    pass


UNIDADES = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove", "dez",
            "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
DEZENAS = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
CENTENAS = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos",
            "seiscentos", "setecentos", "oitocentos", "novecentos"]


def _ate_999(n):
    if n == 0:
        return ""
    if n == 100:
        return "cem"
    partes = []
    c, resto = divmod(n, 100)
    if c:
        partes.append(CENTENAS[c])
    if resto:
        if resto < 20:
            partes.append(UNIDADES[resto])
        else:
            d, u = divmod(resto, 10)
            partes.append(DEZENAS[d] + (f" e {UNIDADES[u]}" if u else ""))
    return " e ".join(partes)


def valor_por_extenso(valor):
    """Porta da macro Valor_Extenso.bas — usada nos contratos."""
    if valor is None:
        return ""
    inteiro = int(valor)
    centavos = int(round((float(valor) - inteiro) * 100))

    if inteiro == 0:
        texto = "zero reais"
    else:
        blocos = []
        milhoes, resto = divmod(inteiro, 1_000_000)
        milhares, unidades = divmod(resto, 1000)
        if milhoes:
            blocos.append(f"{_ate_999(milhoes)} {'milhão' if milhoes == 1 else 'milhões'}")
        if milhares:
            blocos.append("mil" if milhares == 1 else f"{_ate_999(milhares)} mil")
        if unidades:
            blocos.append(_ate_999(unidades))
        texto = ", ".join(blocos) + (" real" if inteiro == 1 else " reais")

    if centavos:
        texto += f" e {_ate_999(centavos)} {'centavo' if centavos == 1 else 'centavos'}"
    return texto


def _substituir_no_paragrafo(paragrafo, variaveis):
    """Troca as variáveis preservando a formatação. O Word quebra o texto em
    "runs", então uma variável pode estar partida entre vários — por isso a
    troca é feita no texto concatenado e devolvida no primeiro run."""
    texto = "".join(run.text for run in paragrafo.runs)
    if "#" not in texto:
        return
    novo = texto
    for chave, valor in variaveis.items():
        novo = novo.replace(chave, "" if valor is None else str(valor))
    if novo != texto and paragrafo.runs:
        paragrafo.runs[0].text = novo
        for run in paragrafo.runs[1:]:
            run.text = ""


def _preencher(caminho_modelo, variaveis):
    doc = Document(str(caminho_modelo))
    for paragrafo in doc.paragraphs:
        _substituir_no_paragrafo(paragrafo, variaveis)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _substituir_no_paragrafo(paragrafo, variaveis)
    for secao in doc.sections:
        for container in (secao.header, secao.footer):
            for paragrafo in container.paragraphs:
                _substituir_no_paragrafo(paragrafo, variaveis)
    return doc


def montar_variaveis(dados):
    """Monta o dicionário `#VARIAVEL` -> valor a partir do contexto do projeto.
    Os nomes seguem as colunas da planilha (abas LINHA BASE e ART)."""
    cliente = dados.get("cliente") or {}
    sistema = dados.get("sistema") or {}
    comercial = dados.get("comercial") or {}
    rt = dados.get("responsavel_tecnico") or {}
    modulos = dados.get("modulos") or {}
    inversor = dados.get("inversor") or {}

    return {
        "#NOME PF": cliente.get("nome"),
        "#cliente": cliente.get("nome"),
        "#CPF": cliente.get("cpf"),
        "#cpf": cliente.get("cpf"),
        "#CEP": cliente.get("cep"),
        "#ENDEREÇO": cliente.get("endereco"),
        "#endereço": cliente.get("endereco"),
        "#CIDADE": cliente.get("cidade"),
        "#cid": cliente.get("cidade"),
        "#uf": cliente.get("uf"),
        "#tel": cliente.get("telefone"),
        "#cel": cliente.get("telefone"),
        "#LATITUDE": cliente.get("latitude"),
        "#lat": cliente.get("latitude"),
        "#LONGITUDE": cliente.get("longitude"),
        "#lon": cliente.get("longitude"),
        "#uc": cliente.get("numero_instalacao"),
        "#N° INSTALAÇÃO": cliente.get("numero_instalacao"),

        "#DATA": dados.get("data"),
        "#data": dados.get("data"),
        "#N° DOC TECSOL": dados.get("doc_tecsol"),
        "#doc": dados.get("doc_tecsol"),
        "#N° CONTRATO": dados.get("codigo_contrato"),
        "#MEMORIAL": dados.get("codigo_memorial"),
        "#N° PROJETO": dados.get("codigo_projeto"),
        "#projeto": dados.get("codigo_projeto"),
        "#deprojeto": dados.get("codigo_desenho"),
        "#N° TRT CFT": dados.get("numero_cft"),
        "#TRT": dados.get("numero_cft"),
        "#PEDIDO": dados.get("numero_pedido"),

        "#FAB MOD": modulos.get("fabricante"),
        "#fabricantepainel": modulos.get("fabricante"),
        "#MOD MOD": modulos.get("modelo"),
        "#modelopainel": modulos.get("modelo"),
        "#POT MOD": modulos.get("potencia_w"),
        "#potenciapainel": modulos.get("potencia_w"),
        "#QTD MOD": modulos.get("quantidade"),
        "#quantidadepainel_": modulos.get("quantidade"),
        "#AREA": modulos.get("area_m2"),

        "#TIPO INV": inversor.get("tipo"),
        "#tipoinversor": inversor.get("tipo"),
        "#FAB INV": inversor.get("fabricante"),
        "#fabricanteinversor": inversor.get("fabricante"),
        "#MOD INV": inversor.get("modelo"),
        "#modeloinversor": inversor.get("modelo"),
        "#QTD INV": inversor.get("quantidade"),
        "#quantidadeinversor": inversor.get("quantidade"),
        "#POT INV": inversor.get("potencia_kw"),
        "#N° MPPT": inversor.get("n_mppt"),

        "#POT kWp": sistema.get("potencia_kwp"),
        "#potenciatotalpainel": sistema.get("potencia_kwp"),
        "#POT kW": sistema.get("potencia_kw"),
        "#potenciatotalinversor": sistema.get("potencia_kw"),
        "#GERACAO": sistema.get("geracao_kwh_mes"),
        "#VALOR CONTA": sistema.get("valor_conta"),

        "#VALOR PROJETO": comercial.get("valor_projeto"),
        "#VALORPROJETOEXT": valor_por_extenso(comercial.get("valor_projeto")),
        "#VALOR KIT": comercial.get("valor_kit"),
        "#VALORKITEXT": valor_por_extenso(comercial.get("valor_kit")),
        "#FORMAPAGAMENTO": comercial.get("forma_pagamento"),
        "#FORNECEDOR": comercial.get("fornecedor"),
        "#ESTRUTURA": comercial.get("estrutura"),

        "#resptecnico": rt.get("nome"),
        "#reg": rt.get("registro"),
        "#projetista": rt.get("projetista") or rt.get("nome"),
    }


def gerar(dados, dir_modelos):
    """Gera os .docx do projeto EM MEMÓRIA. Retorna [(nome_arquivo, bytes)].

    Nada é escrito em disco de propósito: em hospedagem com disco efêmero (ou
    quando o worker roda em outra máquina que a API) um arquivo salvo somem no
    próximo deploy e o link de download quebra. Como o contexto do projeto fica
    guardado no banco, o documento é remontado igual sempre que for baixado.

    Documento sem modelo disponível é pulado (e registrado no log), pra uma
    peça faltante não impedir a geração das outras."""
    dir_modelos = Path(dir_modelos)
    variaveis = montar_variaveis(dados)
    # O nome do arquivo usa o código sem prefixo (2026-01-750-001-A), porque o
    # prefixo de cada documento (MD-, CT-, DC-, RT-) é adicionado abaixo.
    codigo = dados.get("doc_tecsol") or dados.get("codigo_projeto") or "PROJETO"
    nome_cliente = re.sub(r"[^\w\s-]", "", (dados.get("cliente") or {}).get("nome") or "cliente").strip()

    gerados, ausentes = [], []
    for doc in DOCUMENTOS:
        modelo = dir_modelos / MODELOS[doc]
        if not modelo.exists():
            ausentes.append(MODELOS[doc])
            continue
        documento = _preencher(modelo, variaveis)
        buffer = io.BytesIO()
        documento.save(buffer)
        gerados.append((f"{PREFIXOS[doc]}-{codigo} - {nome_cliente}.docx", buffer.getvalue()))

    if ausentes:
        logger.warning("Modelos ausentes em %s: %s", dir_modelos, ", ".join(ausentes))
    if not gerados:
        raise GeracaoError(
            f"Nenhum modelo .docx encontrado em {dir_modelos}. "
            f"Esperado: {', '.join(MODELOS.values())}"
        )
    return gerados
