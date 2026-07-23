"""Geracao dos documentos: valor por extenso, troca das variaveis nos modelos
.docx, regra do cliente unico (nome + CPF) e download em zip."""

import json

import pytest
from docx import Document

from app.documentos_gerador import gerar, montar_variaveis, valor_por_extenso
from app.geracao_documentos import _numero, _texto


@pytest.fixture
def modelos(tmp_path):
    """Modelos .docx minimos com as variaveis que a planilha usa."""
    pasta = tmp_path / "modelos"
    pasta.mkdir()
    conteudos = {
        "memorial.docx": ["Memorial #MEMORIAL", "Cliente: #NOME PF (#CPF)", "Potencia: #POT kWp kWp"],
        "contrato.docx": ["Contrato #N° CONTRATO", "Valor: R$ #VALOR PROJETO (#VALORPROJETOEXT)"],
        "procuracao.docx": ["Outorgante: #NOME PF", "Outorgado: #resptecnico CFT #reg"],
        "rt.docx": ["TRT #TRT", "Projeto #projeto"],
    }
    for nome, linhas in conteudos.items():
        doc = Document()
        for linha in linhas:
            doc.add_paragraph(linha)
        doc.save(str(pasta / nome))
    return pasta


def contexto():
    return {
        "data": "21/07/2026",
        "doc_tecsol": "2026-07-750-001-A",
        "codigo_projeto": "PJ-2026-07-750-001-A",
        "codigo_contrato": "CT-2026-07-750-001-A",
        "codigo_memorial": "MD-2026-07-750-001-A",
        "codigo_desenho": "DE-2026-07-750-001-A",
        "numero_cft": "CFT123",
        "numero_pedido": "07",
        "cliente": {"nome": "Leandro dos Reis", "cpf": "13320018744", "cidade": "Sooretama", "uf": "ES"},
        "sistema": {"potencia_kwp": 8.75, "geracao_kwh_mes": 400},
        "comercial": {"valor_projeto": 21500, "valor_kit": 11800},
        "responsavel_tecnico": {"nome": "William Couto Pereira", "registro": "10694446777"},
        "modulos": {},
        "inversor": {},
    }


class TestValorPorExtenso:
    """Porta da macro Valor_Extenso.bas — vai impresso no contrato."""

    @pytest.mark.parametrize(
        "valor,esperado",
        [
            (0, "zero reais"),
            (1, "um real"),
            (100, "cem reais"),
            (21500, "vinte e um mil, quinhentos reais"),
            (11800, "onze mil, oitocentos reais"),
            (1000000, "um milhão reais"),
        ],
    )
    def test_valores(self, valor, esperado):
        assert valor_por_extenso(valor) == esperado

    def test_com_centavos(self):
        assert valor_por_extenso(1.5) == "um real e cinquenta centavos"

    def test_none_vira_vazio(self):
        assert valor_por_extenso(None) == ""


class TestVariaveis:
    def test_monta_a_partir_do_contexto(self):
        v = montar_variaveis(contexto())
        assert v["#NOME PF"] == "Leandro dos Reis"
        assert v["#CPF"] == "13320018744"
        assert v["#VALOR PROJETO"] == 21500
        assert v["#VALORPROJETOEXT"] == "vinte e um mil, quinhentos reais"

    def test_equipamento_ausente_nao_quebra(self):
        """A Solarz ainda nao expoe equipamentos: as variaveis ficam vazias,
        mas a geracao nao pode falhar por causa disso."""
        v = montar_variaveis(contexto())
        assert v["#FAB MOD"] is None


class TestGeracao:
    """A geracao devolve (nome, bytes) — nada e escrito em disco, pra o
    download funcionar mesmo em servidor com disco efemero."""

    def test_gera_os_quatro_documentos(self, modelos):
        arquivos = gerar(contexto(), modelos)
        assert len(arquivos) == 4
        prefixos = sorted(nome.split("-")[0] for nome, _ in arquivos)
        assert prefixos == ["CT", "DC", "MD", "RT"]

    def test_devolve_bytes_de_docx_valido(self, modelos):
        for _, conteudo in gerar(contexto(), modelos):
            assert isinstance(conteudo, bytes) and len(conteudo) > 0
            assert conteudo[:2] == b"PK"  # .docx e um zip

    def test_nome_do_arquivo_segue_o_padrao_da_engenharia(self, modelos):
        nomes = [nome for nome, _ in gerar(contexto(), modelos)]
        # MD-2026-07-750-001-A - Leandro dos Reis.docx (sem duplicar o "PJ-")
        assert any(n.startswith("MD-2026-07-750-001-A - Leandro dos Reis") for n in nomes)
        assert not any(n.startswith("MD-PJ-") for n in nomes)

    def test_variaveis_sao_substituidas_no_texto(self, modelos):
        import io

        arquivos = dict(gerar(contexto(), modelos))
        nome_contrato = next(n for n in arquivos if n.startswith("CT-"))
        texto = "\n".join(p.text for p in Document(io.BytesIO(arquivos[nome_contrato])).paragraphs)
        assert "#" not in texto
        assert "21500" in texto and "vinte e um mil, quinhentos reais" in texto

    def test_sem_modelo_nenhum_da_erro_claro(self, tmp_path):
        from app.documentos_gerador import GeracaoError

        with pytest.raises(GeracaoError, match="Nenhum modelo"):
            gerar(contexto(), tmp_path / "vazia")

    def test_modelo_faltando_gera_os_outros(self, modelos):
        (modelos / "rt.docx").unlink()
        assert len(gerar(contexto(), modelos)) == 3


class TestLeituraDosCamposDaSolarz:
    def test_numero_aceita_formatos_variados(self):
        assert _numero("400") == 400
        assert _numero("400 kWh") == 400
        assert _numero("1.234,56") == 1234.56
        assert _numero("8.75") == 8.75
        assert _numero(None) is None
        assert _numero("sem numero") is None

    def test_campo_em_json_vira_texto_legivel(self):
        """Forma de Pagamento chega como '["À Vista - Pix"]' na Solarz."""
        assert _texto('["À Vista - Pix"]') == "À Vista - Pix"
        assert _texto('["A", "B"]') == "A, B"
        assert _texto("texto simples") == "texto simples"
        assert _texto(None) is None


class TestClienteUnico:
    """Regra do dono: o trabalho nao e recorrente. Mesmo nome + CPF nao cria
    cliente novo, so pendura outro projeto no registro existente."""

    def test_reaproveita_cliente_com_mesmo_nome_e_cpf(self, db):
        from app.geracao_documentos import _cliente_id

        primeiro = _cliente_id(db, "Joao da Silva", "11122233344")
        segundo = _cliente_id(db, "Joao da Silva", "11122233344")
        db.commit()
        assert primeiro == segundo
        assert db.execute("SELECT COUNT(*) FROM clientes").fetchone()[0] == 1

    def test_cpf_diferente_cria_outro_cliente(self, db):
        from app.geracao_documentos import _cliente_id

        a = _cliente_id(db, "Joao da Silva", "11122233344")
        b = _cliente_id(db, "Joao da Silva", "99988877766")
        db.commit()
        assert a != b


class TestDownload:
    """Os .docx sao remontados do contexto guardado no banco, entao o download
    funciona mesmo depois de um deploy que apague o disco."""

    def _geracao_pronta(self, db, ctx=None, arquivos=None):
        cur = db.execute(
            """INSERT INTO geracoes_documentos
               (solarz_deal_id, numero_pedido, numero_cft, status, contexto, arquivos, projeto_codigo)
               VALUES (?, ?, ?, 'pronto', ?, ?, ?)""",
            (
                253, "07", "CFT123",
                json.dumps(ctx) if ctx is not None else None,
                json.dumps(arquivos or []),
                "PJ-2026-07-750-001-A",
            ),
        )
        db.commit()
        return cur.lastrowid

    def test_zip_traz_todos_os_arquivos(self, admin, db, modelos, app):
        import io
        import zipfile

        app.config["MODELOS_DIR"] = str(modelos)
        nomes = [n for n, _ in gerar(contexto(), modelos)]
        gid = self._geracao_pronta(db, contexto(), nomes)

        resp = admin.get(f"/api/v1/geracao-documentos/{gid}/download")
        assert resp.status_code == 200
        zf = zipfile.ZipFile(io.BytesIO(resp.data))
        assert sorted(zf.namelist()) == sorted(nomes)

    def test_baixar_um_documento(self, admin, db, modelos, app):
        app.config["MODELOS_DIR"] = str(modelos)
        nomes = [n for n, _ in gerar(contexto(), modelos)]
        gid = self._geracao_pronta(db, contexto(), nomes)

        resp = admin.get(f"/api/v1/geracao-documentos/{gid}/arquivos/{nomes[0]}")
        assert resp.status_code == 200
        assert resp.data[:2] == b"PK"

    def test_sem_contexto_guardado_da_404(self, admin, db):
        """Geracao antiga (ou interrompida) nao tem como ser remontada."""
        gid = self._geracao_pronta(db, ctx=None)
        assert admin.get(f"/api/v1/geracao-documentos/{gid}/download").status_code == 404

    def test_download_exige_login(self, client, db):
        gid = self._geracao_pronta(db, contexto(), ["a.docx"])
        assert client.get(f"/api/v1/geracao-documentos/{gid}/download").status_code == 401


class TestNumeroDoPedido:
    def test_comeca_em_01(self, admin):
        assert admin.get("/api/v1/geracao-documentos/proximo-pedido").get_json()["numero_pedido"] == "01"

    def test_sequencial(self, admin, db):
        db.execute(
            "INSERT INTO geracoes_documentos (solarz_deal_id, numero_pedido, numero_cft) VALUES (1, '07', 'x')"
        )
        db.commit()
        assert admin.get("/api/v1/geracao-documentos/proximo-pedido").get_json()["numero_pedido"] == "08"

    def test_gerar_sem_cft_da_400(self, admin):
        resp = admin.post("/api/v1/geracao-documentos", json={"solarz_deal_id": 1, "numero_pedido": "01"})
        assert resp.status_code == 400

    def test_gerar_sem_pedido_da_400(self, admin):
        resp = admin.post("/api/v1/geracao-documentos", json={"solarz_deal_id": 1, "numero_cft": "x"})
        assert resp.status_code == 400
